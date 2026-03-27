"""DOCX Exporter — Convert Markdown CV to a styled .docx document."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def export_docx(markdown_cv: str, output_path: str | Path) -> Path:
    """Convert a Markdown-formatted CV into a styled DOCX file.

    Handles:
    - H1/H2/H3 headings
    - Bold text (**text**)
    - Bullet points (lines starting with -)
    - Regular paragraphs

    Args:
        markdown_cv: The CV content in Markdown.
        output_path: Where to save the .docx file.

    Returns:
        Path to the generated DOCX.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # ─── Global styles ───
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

    # Split content at the --- separator (CV vs justification)
    parts = markdown_cv.split("\n---\n", maxsplit=1)
    cv_content = parts[0]

    for line in cv_content.split("\n"):
        stripped = line.strip()

        if not stripped:
            continue

        # ─── Headings ───
        if stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ─── Bullet points ───
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_text(p, text)

        # ─── Regular text ───
        else:
            p = doc.add_paragraph()
            _add_formatted_text(p, stripped)

    doc.save(str(output_path))
    return output_path


def _add_formatted_text(paragraph: object, text: str) -> None:
    """Add text to a paragraph, handling **bold** markers."""
    # Split on bold markers
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])  # type: ignore[union-attr]
            run.bold = True
        else:
            paragraph.add_run(part)  # type: ignore[union-attr]
